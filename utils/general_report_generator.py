import numbers
import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .globals.constants import string_zero, result_point_14, string_yes
from .globals.report_constants import col_2_values, col_4_values
from .globals.report_constants import rupee_symbol, bo_cs_NA, gstr3b_analysis_NA, gstr2a_analysis_NA, gstr9_NA, \
    ewb_out_analysis_NA, gstr3b_merged_NA, gstr1_merged_NA, gstr1_analysis_NA


async def general_analysis_report_generator(gstin, master_dict):
    print(f"[General report generator] Started execution of method general_analysis_report_generator for: {gstin} ===")
    doc_name = f"{gstin}_GENERAL_REPORT.docx"
    output_path = f"reports/{gstin}/"
    try:
        print(f"Printing master dict: ")
        for outer_key, inner_dict in master_dict.items():
            print(f"{outer_key}:")
            for key, value in inner_dict.items():
                print(f"  {key}: {value}")
            print()  # add empty line between items

        os.makedirs(output_path, exist_ok=True)  # Create dirs if not exist
        output_path = os.path.join(output_path, doc_name)

        gstin_of_taxpayer = master_dict.get('gstr3b_analysis_dict', {}).get('gstin_of_taxpayer', None)
        if gstin_of_taxpayer is None:
            gstin_of_taxpayer = master_dict.get('details_of_taxpayer', {}).get('gstin_of_taxpayer', '')
        legal_name_of_taxpayer = master_dict.get('gstr3b_analysis_dict', {}).get("legal_name_of_taxpayer", "")
        trade_name_of_taxpayer = master_dict.get('gstr3b_analysis_dict', {}).get("trade_name_of_taxpayer", "")

        # Create document
        doc = Document()
        # Set margins (1 inch all sides)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Add Heading
        heading = doc.add_paragraph('General Analysis Report')
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = heading.runs[0]
        run.bold = True
        run.font.size = Pt(14)

        # Add underline using bottom border
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Trick: apply bottom border
        p_border = p._element
        pPr = p_border.get_or_add_pPr()
        borders = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')  # single line
        bottom.set(qn('w:sz'), '8')  # thickness
        bottom.set(qn('w:space'), '1')  # space
        bottom.set(qn('w:color'), 'blue')  # color
        borders.append(bottom)
        pPr.append(borders)

        # Add To Address
        doc.add_paragraph(
            f'\nGSTIN: {gstin_of_taxpayer}'
            f'\nLegal Name: {legal_name_of_taxpayer}'
            f'\nTrade Name: {trade_name_of_taxpayer}')

        # Add table with headers
        table = doc.add_table(rows=2, cols=7)
        table.style = 'Table Grid'
        table.columns[0].width = Inches(0.5)  # Example: Set width of 1st column (Sl. No.) to 0.5 inch
        table.autofit = False

        # First header row
        hdr_cells1 = table.rows[0].cells
        hdr_cells1[0].text = "S. No."
        hdr_cells1[1].text = "Parameter"
        hdr_cells1[2].text = "Quantification of Tax (in ₹)"
        hdr_cells1[6].text = "Legal Provision"

        # Center-align the header text
        for cell in hdr_cells1:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Merge Quantification of Tax across 4 columns
        tax_header_cell = hdr_cells1[2]
        for i in range(3, 6):
            tax_header_cell.merge(hdr_cells1[i])

        # Second header row (sub-columns)
        hdr_cells2 = table.rows[1].cells
        hdr_cells2[0].text = ''  # empty because merged vertically
        hdr_cells2[1].text = ''  # empty because merged vertically
        hdr_cells2[2].text = 'IGST'
        hdr_cells2[3].text = 'CGST'
        hdr_cells2[4].text = 'SGST'
        hdr_cells2[5].text = 'Cess'
        hdr_cells2[6].text = ''  # empty because merged vertically

        for cell in hdr_cells2:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Merge vertically for first 2 columns and last column
        hdr_cells1[0].merge(table.rows[1].cells[0])
        hdr_cells1[1].merge(table.rows[1].cells[1])
        hdr_cells1[6].merge(table.rows[1].cells[6])

        # Force column widths
        for row in table.rows:
            row.cells[0].width = Inches(0.3)  # Sl. No.
            row.cells[1].width = Inches(1.7)  # Parameter
            row.cells[2].width = Inches(1)  # IGST
            row.cells[3].width = Inches(1)  # CGST
            row.cells[4].width = Inches(1)  # SGST
            row.cells[5].width = Inches(1)  # Cess
            row.cells[6].width = Inches(1)  # Legal Provisions Violated

        # Populate columns 1, 2 & 4
        num_rows = len(col_2_values)
        row_number = 1  # Start from 1
        i = 0
        while i < num_rows:
            row_cells = table.add_row().cells
            if i == 5:
                row_cells[0].text = "5.1"  # Special row labeled '5.1'
            else:
                row_cells[0].text = str(row_number)  # Normal row numbering
                row_number += 1  # Only increment on standard rows
            row_cells[1].text = col_2_values[i]
            row_cells[6].text = col_4_values[i]
            i += 1  # Always increment data index

        # Populate columns based on master_dict values
        table = doc.tables[0]  # Assuming the table is the first one in the document

        row_pos = 2  # We start from 2 as first 2 (0 & 1) are header rows
        # 1. Row 1 cell 3
        igst = master_dict.get('bo_comparison_summary_dict', {}).get('result_point_1_IGST', None)
        cgst = master_dict.get('bo_comparison_summary_dict', {}).get('result_point_1_CGST', None)
        sgst = master_dict.get('bo_comparison_summary_dict', {}).get('result_point_1_SGST', None)
        cess = master_dict.get('bo_comparison_summary_dict', {}).get('result_point_1_CESS', None)
        if isinstance(igst, numbers.Number):
            if igst < 0:  # populate only if igst is -ve
                table.rows[row_pos].cells[2].text = rupee_symbol + str(abs(igst))
            else:
                table.rows[row_pos].cells[2].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 1 cell 3 igst is +ve or 0.00: {igst}, replacing by 0 instead.")
        elif igst is None:
            print(f"Row 1 cell 3: igst is not proper: igst = {igst}")
            table.rows[row_pos].cells[2].text = bo_cs_NA
        #  cgst
        if isinstance(cgst, numbers.Number):
            if cgst < 0:  # populate only if cgst is -ve
                table.rows[row_pos].cells[3].text = rupee_symbol + str(abs(cgst))
            else:
                table.rows[row_pos].cells[3].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 1 cell 4 cgst is +ve or 0.00: {cgst}, replacing by 0 instead.")
        elif cgst is None:
            print(f"Row 1 cell 4: cgst is not proper: cgst = {cgst}")
            table.rows[row_pos].cells[3].text = bo_cs_NA
        #  sgst
        if isinstance(sgst, numbers.Number):
            if sgst < 0:  # populate only if sgst is -ve
                table.rows[row_pos].cells[4].text = rupee_symbol + str(abs(sgst))
            else:
                table.rows[row_pos].cells[4].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 1 cell 5 sgst is +ve or 0.00: {sgst}, replacing by 0 instead.")
        elif sgst is None:
            print(f"Row 1 cell 5: sgst is not proper: sgst = {sgst}")
            table.rows[row_pos].cells[4].text = bo_cs_NA
        #  cess
        if isinstance(cess, numbers.Number):
            if cess < 0:  # populate only if cess is -ve
                table.rows[row_pos].cells[5].text = rupee_symbol + str(abs(cess))
            else:
                table.rows[row_pos].cells[5].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 1 cell 6 cess is +ve or 0.00: {cess}, replacing by 0 instead.")
        elif cess is None:
            print(f"Row 1 cell 6: cess is not proper: cess = {cess}")
            table.rows[row_pos].cells[5].text = bo_cs_NA
        row_pos += 1

        # 2. Row 2 cell 3
        igst = master_dict.get('bo_comparison_summary_dict', {}).get('result_point_2_IGST', None)
        cgst = master_dict.get('bo_comparison_summary_dict', {}).get('result_point_2_CGST', None)
        sgst = master_dict.get('bo_comparison_summary_dict', {}).get('result_point_2_SGST', None)
        cess = master_dict.get('bo_comparison_summary_dict', {}).get('result_point_2_CESS', None)
        if isinstance(igst, numbers.Number):
            if igst < 0:  # populate only if igst is -ve
                table.rows[row_pos].cells[2].text = rupee_symbol + str(abs(igst))
            else:
                table.rows[row_pos].cells[2].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 2 cell 3 igst is +ve or 0.00: {igst}, replacing by 0 instead.")
        elif igst is None:
            print(f"Row 2 cell 3: igst is not proper: igst = {igst}")
            table.rows[row_pos].cells[2].text = bo_cs_NA
        #  cgst
        if isinstance(cgst, numbers.Number):
            if cgst < 0:  # populate only if cgst is -ve
                table.rows[row_pos].cells[3].text = rupee_symbol + str(abs(cgst))
            else:
                table.rows[row_pos].cells[3].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 2 cell 4 cgst is +ve or 0.00: {cgst}, replacing by 0 instead.")
        elif cgst is None:
            print(f"Row 2 cell 4: cgst is not proper: cgst = {cgst}")
            table.rows[row_pos].cells[3].text = bo_cs_NA
        #  sgst
        if isinstance(sgst, numbers.Number):
            if sgst < 0:  # populate only if sgst is -ve
                table.rows[row_pos].cells[4].text = rupee_symbol + str(abs(sgst))
            else:
                table.rows[row_pos].cells[4].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 2 cell 5 sgst is +ve or 0.00: {sgst}, replacing by 0 instead.")
        elif sgst is None:
            print(f"Row 2 cell 5: sgst is not proper: sgst = {sgst}")
            table.rows[row_pos].cells[4].text = bo_cs_NA
        #  cess
        if isinstance(cess, numbers.Number):
            if cess < 0:  # populate only if cess is -ve
                table.rows[row_pos].cells[5].text = rupee_symbol + str(abs(cess))
            else:
                table.rows[row_pos].cells[5].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 2 cell 6 cess is +ve or 0.00: {cess}, replacing by 0 instead.")
        elif cess is None:
            print(f"Row 2 cell 6: cess is not proper: cess = {cess}")
            table.rows[row_pos].cells[5].text = bo_cs_NA
        row_pos += 1

        # 3. Row 3 cell 3
        igst = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_3_IGST', None)
        cgst = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_3_CGST', None)
        sgst = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_3_SGST', None)
        cess = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_3_CESS', None)
        if igst is not None:
            table.rows[row_pos].cells[2].text = rupee_symbol + str(igst)
        else:
            table.rows[row_pos].cells[2].text = gstr3b_analysis_NA
        if cgst is not None:
            table.rows[row_pos].cells[3].text = rupee_symbol + str(cgst)
        else:
            table.rows[row_pos].cells[3].text = gstr3b_analysis_NA
        if sgst is not None:
            table.rows[row_pos].cells[4].text = rupee_symbol + str(sgst)
        else:
            table.rows[row_pos].cells[4].text = gstr3b_analysis_NA
        if cess is not None:
            table.rows[row_pos].cells[5].text = rupee_symbol + str(cess)
        else:
            table.rows[row_pos].cells[5].text = gstr3b_analysis_NA
        row_pos += 1

        # 4. Row 4 cell 3
        igst = master_dict.get('bo_comparison_summary_dict', {}).get('result_point_4_IGST', None)
        cgst = master_dict.get('bo_comparison_summary_dict', {}).get('result_point_4_CGST', None)
        sgst = master_dict.get('bo_comparison_summary_dict', {}).get('result_point_4_SGST', None)
        cess = master_dict.get('bo_comparison_summary_dict', {}).get('result_point_4_CESS', None)
        if isinstance(igst, numbers.Number):
            if igst > 0:  # populate only if value is +ve
                table.rows[row_pos].cells[2].text = rupee_symbol + str(igst)
            else:
                table.rows[row_pos].cells[2].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 4 cell 3 igst is -ve or 0.00: {igst}, replacing by 0 instead.")
        elif igst is None:
            print(f"Row 4 cell 3: igst is not proper: value = {igst}")
            table.rows[row_pos].cells[2].text = bo_cs_NA
        if isinstance(cgst, numbers.Number):
            if cgst > 0:  # populate only if value is +ve
                table.rows[row_pos].cells[3].text = rupee_symbol + str(cgst)
            else:
                table.rows[row_pos].cells[3].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 4 cell 4 cgst is -ve or 0.00: {cgst}, replacing by 0 instead.")
        elif cgst is None:
            print(f"Row 4 cell 4: cgst is not proper: value = {cgst}")
            table.rows[row_pos].cells[3].text = bo_cs_NA
        if isinstance(sgst, numbers.Number):
            if sgst > 0:  # populate only if value is +ve
                table.rows[row_pos].cells[4].text = rupee_symbol + str(sgst)
            else:
                table.rows[row_pos].cells[4].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 4 cell 5 sgst is -ve or 0.00: {sgst}, replacing by 0 instead.")
        elif sgst is None:
            print(f"Row 4 cell 5: sgst is not proper: value = {sgst}")
            table.rows[row_pos].cells[4].text = bo_cs_NA
        if isinstance(cess, numbers.Number):
            if cess > 0:  # populate only if value is +ve
                table.rows[row_pos].cells[5].text = rupee_symbol + str(cess)
            else:
                table.rows[row_pos].cells[5].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 4 cell 6 cess is -ve or 0.00: {cess}, replacing by 0 instead.")
        elif cess is None:
            print(f"Row 4 cell 6: cess is not proper: value = {cess}")
            table.rows[row_pos].cells[5].text = bo_cs_NA
        row_pos += 1

        # 5. Row 5 cell 3
        igst = master_dict.get('gstr2a_analysis_dict', {}).get('result_point_5_IGST', None)
        cgst = master_dict.get('gstr2a_analysis_dict', {}).get('result_point_5_CGST', None)
        sgst = master_dict.get('gstr2a_analysis_dict', {}).get('result_point_5_SGST', None)
        cess = master_dict.get('gstr2a_analysis_dict', {}).get('result_point_5_CESS', None)
        if igst is not None:
            table.rows[row_pos].cells[2].text = rupee_symbol + str(igst)
        else:
            table.rows[row_pos].cells[2].text = gstr2a_analysis_NA
        if cgst is not None:
            table.rows[row_pos].cells[3].text = rupee_symbol + str(cgst)
        else:
            table.rows[row_pos].cells[3].text = gstr2a_analysis_NA
        if sgst is not None:
            table.rows[row_pos].cells[4].text = rupee_symbol + str(sgst)
        else:
            table.rows[row_pos].cells[4].text = gstr2a_analysis_NA
        if cess is not None:
            table.rows[row_pos].cells[5].text = rupee_symbol + str(cess)
        else:
            table.rows[row_pos].cells[5].text = gstr2a_analysis_NA
        row_pos += 1

        # 5.1 Row 5.1 cell 3
        igst = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_5_1_IGST', None)
        cgst = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_5_1_CGST', None)
        sgst = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_5_1_SGST', None)
        cess = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_5_1_CESS', None)
        if isinstance(igst, numbers.Number):
            if igst < 0:
                table.rows[row_pos].cells[2].text = str(abs(igst))
            else:
                print(f"Not accounted: Row 5.1 cell 3 IGST value is +ve or 0.00: {igst}, replacing by 0 instead.")
                table.rows[row_pos].cells[2].text = rupee_symbol + string_zero
        elif igst is None:
            table.rows[row_pos].cells[2].text = gstr9_NA
            print(f"Row 5.1 cell 3: IGST value is not proper: value = {igst}")

        if isinstance(cgst, numbers.Number):
            if cgst < 0:
                table.rows[row_pos].cells[3].text = str(abs(cgst))
            else:
                print(f"Not accounted: Row 5.1 cell 4 CGST value is +ve or 0.00: {cgst}, replacing by 0 instead.")
                table.rows[row_pos].cells[3].text = rupee_symbol + string_zero
        elif cgst is None:
            table.rows[row_pos].cells[3].text = gstr9_NA
            print(f"Row 5.1 cell 4: CGST value is not proper: value = {cgst}")

        if isinstance(sgst, numbers.Number):
            if sgst < 0:
                table.rows[row_pos].cells[4].text = str(abs(sgst))
            else:
                print(f"Not accounted: Row 5.1 cell 5 SGST value is +ve or 0.00: {sgst}, replacing by 0 instead.")
                table.rows[row_pos].cells[4].text = rupee_symbol + string_zero
        elif sgst is None:
            table.rows[row_pos].cells[4].text = gstr9_NA
            print(f"Row 5.1 cell 5: SGST value is not proper: value = {sgst}")

        if isinstance(cess, numbers.Number):
            if cess < 0:
                table.rows[row_pos].cells[5].text = rupee_symbol + str(abs(cess))
            else:
                print(f"Not accounted: Row 5.1 cell 6 CESS value is +ve or 0.00: {cess}, replacing by 0 instead.")
                table.rows[row_pos].cells[5].text = rupee_symbol + string_zero
        elif cess is None:
            table.rows[row_pos].cells[5].text = gstr9_NA
            print(f"Row 5.1 cell 6: CESS value is not proper: value = {cess}")
        row_pos += 1

        # 6. Row 6 cell 3
        value_1 = master_dict.get('bo_comparison_summary_dict', {}).get('result_point_6', None)
        value_2 = master_dict.get('ewb_out_analysis_dict', {}).get('result_point_6', None)
        if value_1 is not None and value_2 is not None:
            if isinstance(value_1, numbers.Number) and isinstance(value_2, numbers.Number):
                difference = value_1 - value_2
                if difference < 0:  # populate only if value is -ve
                    table.rows[row_pos].cells[2].text = rupee_symbol + str(abs(difference))
                else:
                    table.rows[row_pos].cells[2].text = rupee_symbol + string_zero
                    print(
                        f"Not accounted: Row 6 cell 3 value is +ve or 0.00: {difference}, replacing by 0 instead.")
        elif value_1 is None:
            table.rows[row_pos].cells[2].text = bo_cs_NA
            print(
                f"Row 6 cell 3: Either one or both of the values is not proper. value1 = {value_1}, value2 = {value_2}")
        elif value_2 is None:
            table.rows[row_pos].cells[2].text = ewb_out_analysis_NA
            print(
                f"Row 6 cell 3: Either one or both of the values is not proper. value1 = {value_1}, value2 = {value_2}")
        row_pos += 1

        # 7.  Row 7 cell 3
        igst = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_7_IGST', None)
        cgst = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_7_CGST', None)
        sgst = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_7_SGST', None)
        cess = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_7_CESS', None)
        if igst is not None:
            table.rows[row_pos].cells[2].text = rupee_symbol + str(igst)
        else:
            table.rows[row_pos].cells[2].text = gstr3b_analysis_NA
        if cgst is not None:
            table.rows[row_pos].cells[3].text = rupee_symbol + str(cgst)
        else:
            table.rows[row_pos].cells[3].text = gstr3b_analysis_NA
        if sgst is not None:
            table.rows[row_pos].cells[4].text = rupee_symbol + str(sgst)
        else:
            table.rows[row_pos].cells[4].text = gstr3b_analysis_NA
        if cess is not None:
            table.rows[row_pos].cells[5].text = rupee_symbol + str(cess)
        else:
            table.rows[row_pos].cells[5].text = gstr3b_analysis_NA
        row_pos += 1

        # 8.  Row 8 cell 3
        igst = master_dict.get('gstr3b_merged_dict', {}).get('result_point_8_IGST', None)
        cgst = master_dict.get('gstr3b_merged_dict', {}).get('result_point_8_CGST', None)
        sgst = master_dict.get('gstr3b_merged_dict', {}).get('result_point_8_SGST', None)
        cess = master_dict.get('gstr3b_merged_dict', {}).get('result_point_8_CESS', None)
        if igst is not None:
            table.rows[row_pos].cells[2].text = rupee_symbol + str(igst)
        else:
            table.rows[row_pos].cells[2].text = gstr3b_merged_NA
        if cgst is not None:
            table.rows[row_pos].cells[3].text = rupee_symbol + str(cgst)
        else:
            table.rows[row_pos].cells[3].text = gstr3b_merged_NA
        if sgst is not None:
            table.rows[row_pos].cells[4].text = rupee_symbol + str(sgst)
        else:
            table.rows[row_pos].cells[4].text = gstr3b_merged_NA
        if cess is not None:
            table.rows[row_pos].cells[5].text = rupee_symbol + str(cess)
        else:
            table.rows[row_pos].cells[5].text = gstr3b_merged_NA
        row_pos += 1

        # 9. Row 9 cell 3
        igst = master_dict.get('gstr3b_merged_dict', {}).get('result_point_9_IGST', None)
        cgst = master_dict.get('gstr3b_merged_dict', {}).get('result_point_9_CGST', None)
        sgst = master_dict.get('gstr3b_merged_dict', {}).get('result_point_9_SGST', None)
        cess = master_dict.get('gstr3b_merged_dict', {}).get('result_point_9_CESS', None)
        if igst is not None:
            table.rows[row_pos].cells[2].text = rupee_symbol + str(igst)
        else:
            table.rows[row_pos].cells[2].text = gstr3b_merged_NA
        if cgst is not None:
            table.rows[row_pos].cells[3].text = rupee_symbol + str(cgst)
        else:
            table.rows[row_pos].cells[3].text = gstr3b_merged_NA
        if sgst is not None:
            table.rows[row_pos].cells[4].text = rupee_symbol + str(sgst)
        else:
            table.rows[row_pos].cells[4].text = gstr3b_merged_NA
        if cess is not None:
            table.rows[row_pos].cells[5].text = rupee_symbol + str(cess)
        else:
            table.rows[row_pos].cells[5].text = gstr3b_merged_NA
        row_pos += 1

        # 10. Row 10 cell 3
        igst = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_10_IGST', None)
        cgst = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_10_CGST', None)
        sgst = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_10_SGST', None)
        cess = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_10_CESS', None)
        if isinstance(igst, numbers.Number):
            if igst > 0:
                table.rows[row_pos].cells[2].text = f"{rupee_symbol}{igst} as applicable."
            else:
                print(f"Not accounted: Row 10 cell 3 IGST value is -ve or 0.00: {igst}, replacing by 0 instead.")
                table.rows[row_pos].cells[2].text = rupee_symbol + string_zero
        elif igst is None:
            table.rows[row_pos].cells[2].text = gstr9_NA
            print(f"Row 10 cell 3: IGST value is not proper: value = {igst}")

        if isinstance(cgst, numbers.Number):
            if cgst > 0:
                table.rows[row_pos].cells[3].text = f"{rupee_symbol}{cgst} as applicable."
            else:
                print(f"Not accounted: Row 10 cell 4 CGST value is -ve or 0.00: {cgst}, replacing by 0 instead.")
                table.rows[row_pos].cells[3].text = rupee_symbol + string_zero
        elif cgst is None:
            table.rows[row_pos].cells[3].text = gstr9_NA
            print(f"Row 10 cell 4: CGST value is not proper: value = {cgst}")

        if isinstance(sgst, numbers.Number):
            if sgst > 0:
                table.rows[row_pos].cells[4].text = f"{rupee_symbol}{sgst} as applicable."
            else:
                print(f"Not accounted: Row 10 cell 5 SGST value is -ve or 0.00: {sgst}, replacing by 0 instead.")
                table.rows[row_pos].cells[4].text = rupee_symbol + string_zero
        elif sgst is None:
            table.rows[row_pos].cells[4].text = gstr9_NA
            print(f"Row 10 cell 5: SGST value is not proper: value = {sgst}")

        if isinstance(cess, numbers.Number):
            if cess > 0:
                table.rows[row_pos].cells[5].text = f"{rupee_symbol}{cess} as applicable."
            else:
                print(f"Not accounted: Row 10 cell 6 CESS value is -ve or 0.00: {cess}, replacing by 0 instead.")
                table.rows[row_pos].cells[5].text = rupee_symbol + string_zero
        elif cess is None:
            table.rows[row_pos].cells[5].text = gstr9_NA
            print(f"Row 10 cell 6: CESS value is not proper: value = {cess}")
        row_pos += 1

        # 11. Row 11 cell 3
        igst = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_11_IGST', None)
        cgst = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_11_CGST', None)
        sgst = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_11_SGST', None)
        cess = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_11_CESS', None)
        if isinstance(igst, numbers.Number):
            if igst > 0:
                table.rows[row_pos].cells[2].text = rupee_symbol + str(igst) + " as applicable."
            else:
                print(f"Not accounted: Row 11 cell 3 IGST value is -ve or 0.00: {igst}, replacing by 0 instead.")
                table.rows[row_pos].cells[2].text = rupee_symbol + string_zero
        elif igst is None:
            table.rows[row_pos].cells[2].text = gstr9_NA
            print(f"Row 11 cell 3: IGST value is not proper: value = {igst}")

        if isinstance(cgst, numbers.Number):
            if cgst > 0:
                table.rows[row_pos].cells[3].text = rupee_symbol + str(cgst) + " as applicable."
            else:
                print(f"Not accounted: Row 11 cell 4 CGST value is -ve or 0.00: {cgst}, replacing by 0 instead.")
                table.rows[row_pos].cells[3].text = rupee_symbol + string_zero
        elif cgst is None:
            table.rows[row_pos].cells[3].text = gstr9_NA
            print(f"Row 11 cell 4: CGST value is not proper: value = {cgst}")

        if isinstance(sgst, numbers.Number):
            if sgst > 0:
                table.rows[row_pos].cells[4].text = rupee_symbol + str(sgst) + " as applicable."
            else:
                print(f"Not accounted: Row 11 cell 5 SGST value is -ve or 0.00: {sgst}, replacing by 0 instead.")
                table.rows[row_pos].cells[4].text = rupee_symbol + string_zero
        elif sgst is None:
            table.rows[row_pos].cells[4].text = gstr9_NA
            print(f"Row 11 cell 5: SGST values is not proper: value = {sgst}")

        if isinstance(cess, numbers.Number):
            if cess > 0:
                table.rows[row_pos].cells[5].text = rupee_symbol + str(cess) + " as applicable."
            else:
                print(f"Not accounted: Row 11 cell 6 CESS value is -ve or 0.00: {cess}, replacing by 0 instead.")
                table.rows[row_pos].cells[5].text = rupee_symbol + string_zero
        elif cess is None:
            table.rows[row_pos].cells[5].text = gstr9_NA
            print(f"Row 11 cell 6: CESS value is not proper: value = {cess}")
        row_pos += 1

        # 12. Row 12 cell 3
        value_1 = master_dict.get('gstr3b_merged_dict', {}).get('result_point_12_total_late_fee_gstr3b', None)
        value_2 = master_dict.get('gstr1_merged_dict', {}).get('result_point_12_total_late_fee_gstr1', None)
        value_3 = master_dict.get('gstr3b_merged_dict', {}).get('result_point_12_late_fee_paid_in_cash', None)
        if isinstance(value_1, numbers.Number) and isinstance(value_2, numbers.Number):
            total_late_fee = value_1 + value_2
            if isinstance(value_3, numbers.Number):
                total_late_fee = total_late_fee - value_3
                if total_late_fee > 0:
                    # Divide total late fee into two parts: CGST & SGST
                    table.rows[row_pos].cells[3].text = rupee_symbol + str(round(total_late_fee / 2, 2))
                    table.rows[row_pos].cells[4].text = rupee_symbol + str(round(total_late_fee / 2, 2))
                else:
                    print(f"Not accounted: Row 12 cell 3 value is -ve or 0.00: {total_late_fee}")
                    table.rows[row_pos].cells[3].text = rupee_symbol + string_zero
                    table.rows[row_pos].cells[4].text = rupee_symbol + string_zero
            else:
                table.rows[row_pos].cells[3].text = rupee_symbol + str(round(total_late_fee / 2, 2))
                table.rows[row_pos].cells[4].text = rupee_symbol + str(round(total_late_fee / 2, 2))
                print(f"Row 12 cell 3 Late fee paid in cash is : {value_3}")
        elif value_1 is None or value_3 is None:
            table.rows[row_pos].cells[3].text = gstr3b_merged_NA
        elif value_2 is None:
            table.rows[row_pos].cells[4].text = gstr1_merged_NA
        row_pos += 1

        # 13. Row 13 cell 3 - new addition GSTR-9 late fee
        value = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_13', None)
        if value is not None:
            table.rows[row_pos].cells[3].text = rupee_symbol + str(round(value / 2, 2))
            table.rows[row_pos].cells[4].text = rupee_symbol + str(round(value / 2, 2))
        else:
            table.rows[row_pos].cells[3].text = gstr9_NA
            table.rows[row_pos].cells[4].text = gstr9_NA
        row_pos += 1

        # 14. Row 14 cell 3
        value = master_dict.get('gstr1_analysis_dict', {}).get(result_point_14, None)
        cell = table.rows[row_pos].cells[2]
        paragraph = cell.paragraphs[0]
        paragraph.clear()
        if value is None:
            paragraph.add_run(gstr1_analysis_NA)
        elif string_yes == value:
            paragraph.add_run("Differential rate of tax paid by the taxpayer on the same HSN.")
            run2 = paragraph.add_run(" Please check liability as per applicable rate.")
            run2.font.color.rgb = RGBColor(255, 0, 0)  # Red text
        else:
            paragraph.add_run("No differential tax paid on the same HSN found.")
            print("Row 14: No differential tax paid on the same HSN found.")
        row_pos += 1

        # 15. Row 15 cell 3
        value_1 = master_dict.get('gstr2a_analysis_dict', {}).get('result_point_15', None)
        value_2 = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_15', None)
        if isinstance(value_1, numbers.Number) and isinstance(value_2, numbers.Number):
            if value_1 - value_2 > 0:  # populate only if value is +ve
                table.rows[row_pos].cells[2].text = rupee_symbol + str(abs(value_1 - value_2))
            else:
                table.rows[row_pos].cells[2].text = rupee_symbol + string_zero
                print(
                    f"Not accounted: Row 15 cell 3 value is -ve or 0.00: {value_1 - value_2}, replacing by 0 instead.")
        elif value_1 is None:
            table.rows[row_pos].cells[2].text = gstr2a_analysis_NA
            print(
                f"Row 15 cell 3: Either one or both of the values is not proper. value1: {value_1}, value2 = {value_2}")
        elif value_2 is None:
            table.rows[row_pos].cells[2].text = gstr3b_analysis_NA
            print(
                f"Row 15 cell 3: Either one or both of the values is not proper. value1: {value_1}, value2 = {value_2}")
        row_pos += 1

        # 16. Row 16 cell 3
        value_1 = master_dict.get('gstr2a_analysis_dict', {}).get('result_point_16', None)
        value_2 = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_16', None)
        if isinstance(value_1, numbers.Number) and isinstance(value_2, numbers.Number):
            difference = value_1 - value_2
            if difference > 0:  # populate only if value is +ve
                table.rows[row_pos].cells[2].text = rupee_symbol + str(value_1 - value_2)
            else:
                table.rows[row_pos].cells[2].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 16 cell 3 value is -ve or 0.00: {difference}, replacing by 0 instead.")
        elif value_1 is None:
            table.rows[row_pos].cells[2].text = gstr2a_analysis_NA
            print(
                f"Row 16 cell 3: Either one or both of the values is not proper. value1: {value_1}, value2 = {value_2}")
        elif value_2 is None:
            table.rows[row_pos].cells[2].text = gstr3b_analysis_NA
            print(
                f"Row 16 cell 3: Either one or both of the values is not proper. value1: {value_1}, value2 = {value_2}")
        row_pos += 1

        # 17. Row 17 cell 3
        value = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_17', None)
        if value is None:
            table.rows[row_pos].cells[2].text = gstr9_NA
        else:
            table.rows[row_pos].cells[2].text = rupee_symbol + str(value)
        row_pos += 1

        # 18. Row 18 cell 3
        igst = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_18_IGST', None)
        cgst = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_18_CGST', None)
        sgst = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_18_SGST', None)
        cess = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_18_CESS', None)
        if isinstance(igst, numbers.Number):
            if igst < 0:  # populate only if value is -ve
                table.rows[row_pos].cells[2].text = rupee_symbol + str(abs(igst))
            else:
                table.rows[row_pos].cells[2].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 18 cell 3 IGST is +ve or 0.00: {igst}, replacing by 0 instead.")
        elif igst is None:
            table.rows[row_pos].cells[2].text = gstr9_NA
            print(f"Row 18 cell 3: IGST value is not proper.  value = {igst}")

        if isinstance(cgst, numbers.Number):
            if cgst < 0:  # populate only if value is -ve
                table.rows[row_pos].cells[3].text = rupee_symbol + str(abs(cgst))
            else:
                table.rows[row_pos].cells[3].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 18 cell 4 CGST is +ve or 0.00: {cgst}, replacing by 0 instead.")
        elif cgst is None:
            table.rows[row_pos].cells[3].text = gstr9_NA
            print(f"Row 18 cell 4: CGST value is not proper.  value = {cgst}")

        if isinstance(sgst, numbers.Number):
            if sgst < 0:  # populate only if value is -ve
                table.rows[row_pos].cells[4].text = rupee_symbol + str(abs(sgst))
            else:
                table.rows[row_pos].cells[4].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 18 cell 5 SGST is +ve or 0.00: {sgst}, replacing by 0 instead.")
        elif sgst is None:
            table.rows[row_pos].cells[4].text = gstr9_NA
            print(f"Row 18 cell 5: SGST value is not proper.  value = {sgst}")

        if isinstance(cess, numbers.Number):
            if cess < 0:  # populate only if value is -ve
                table.rows[row_pos].cells[5].text = rupee_symbol + str(abs(cess))
            else:
                table.rows[row_pos].cells[5].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 18 cell 6 CESS is +ve or 0.00: {cess}, replacing by 0 instead.")
        elif cess is None:
            table.rows[row_pos].cells[5].text = gstr9_NA
            print(f"Row 18 cell 6: CESS value is not proper.  value = {cess}")
        row_pos += 1

        # 19. Row 19 cell 3
        igst = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_19_IGST', None)
        cgst = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_19_CGST', None)
        sgst = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_19_SGST', None)
        cess = master_dict.get('gstr9_Vs_3b_analysis_dict', {}).get('result_point_19_CESS', None)
        if isinstance(igst, numbers.Number):
            if igst > 0:
                table.rows[row_pos].cells[2].text = rupee_symbol + str(igst)
            else:  # if igst is -ve put 0
                table.rows[row_pos].cells[2].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 19 cell 3 IGST is -ve or 0.00, replacing by 0 instead. {igst}")
        elif igst is None:
            table.rows[row_pos].cells[2].text = gstr9_NA
            print(f"Row 19 cell 3: IGST is not proper: value = {igst}")

        if isinstance(cgst, numbers.Number):
            if cgst > 0:
                table.rows[row_pos].cells[3].text = rupee_symbol + str(cgst)
            else:  # if cgst is -ve put 0
                table.rows[row_pos].cells[3].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 19 cell 4 CGST is -ve or 0.00, replacing by 0 instead. {cgst}")
        elif cgst is None:
            table.rows[row_pos].cells[3].text = gstr9_NA
            print(f"Row 19 cell 4: CGST is not proper: value = {cgst}")

        if isinstance(sgst, numbers.Number):
            if sgst > 0:
                table.rows[row_pos].cells[4].text = rupee_symbol + str(sgst)
            else:  # if sgst is -ve put 0
                table.rows[row_pos].cells[4].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 19 cell 5 SGST is -ve or 0.00, replacing by 0 instead. {sgst}")
        elif sgst is None:
            table.rows[row_pos].cells[4].text = gstr9_NA
            print(f"Row 19 cell 5: SGST is not proper: value = {sgst}")

        if isinstance(cess, numbers.Number):
            if cess > 0:
                table.rows[row_pos].cells[5].text = rupee_symbol + str(cess)
            else:  # if cess is -ve put 0
                table.rows[row_pos].cells[5].text = rupee_symbol + string_zero
                print(f"Not accounted: Row 19 cell 6 CESS is -ve or 0.00, replacing by 0 instead. {cess}")
        elif cess is None:
            table.rows[row_pos].cells[5].text = gstr9_NA
            print(f"Row 19 cell 6: CESS is not proper: value = {cess}")
        row_pos += 1

        # 20. Row 20 cell 3
        igst1 = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_20_IGST', None)
        cgst1 = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_20_CGST', None)
        sgst1 = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_20_SGST', None)
        cess1 = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_20_CESS', None)
        igst2 = master_dict.get('gstr2a_analysis_dict', {}).get('result_point_20_IGST', None)
        cgst2 = master_dict.get('gstr2a_analysis_dict', {}).get('result_point_20_CGST', None)
        sgst2 = master_dict.get('gstr2a_analysis_dict', {}).get('result_point_20_SGST', None)
        cess2 = master_dict.get('gstr2a_analysis_dict', {}).get('result_point_20_CESS', None)
        if isinstance(igst1, numbers.Number) and isinstance(igst2, numbers.Number):
            difference = igst1 - igst2
            if difference > 0:  # populate if value is +ve
                table.rows[row_pos].cells[2].text = rupee_symbol + str(difference)
            else:
                print(f"Not accounted: Row 20 cell 3 value is -ve or 0.00: {difference}, replacing by 0 instead.")
                table.rows[row_pos].cells[2].text = rupee_symbol + string_zero
        elif igst1 is None:
            table.rows[row_pos].cells[2].text = gstr3b_analysis_NA
            print(
                f"Row 20 cell 3: Either one or both of the IGST values is not proper- igst1: {igst1}, igst2 = {igst2}")
        elif igst2 is None:
            table.rows[row_pos].cells[2].text = gstr2a_analysis_NA
            print(
                f"Row 20 cell 3: Either one or both of the IGST values is not proper- igst1: {igst1}, igst2 = {igst2}")

        if isinstance(cgst1, numbers.Number) and isinstance(cgst2, numbers.Number):
            difference = cgst1 - cgst2
            if difference > 0:  # populate if value is +ve
                table.rows[row_pos].cells[3].text = rupee_symbol + str(difference)
            else:
                print(f"Not accounted: Row 20 cell 4 value is -ve or 0.00: {difference}, replacing by 0 instead.")
                table.rows[row_pos].cells[3].text = rupee_symbol + string_zero
        elif cgst1 is None:
            table.rows[row_pos].cells[3].text = gstr3b_analysis_NA
            print(
                f"Row 20 cell 4: Either one or both of the CGST values is not proper- cgst1: {cgst1}, cgst2 = {cgst2}")
        elif cgst2 is None:
            table.rows[row_pos].cells[3].text = gstr2a_analysis_NA
            print(
                f"Row 20 cell 4: Either one or both of the CGST values is not proper- cgst1: {cgst1}, cgst2 = {cgst2}")

        if isinstance(sgst1, numbers.Number) and isinstance(sgst2, numbers.Number):
            difference = sgst1 - sgst2
            if difference > 0:  # populate if value is +ve
                table.rows[row_pos].cells[4].text = rupee_symbol + str(difference)
            else:
                print(f"Not accounted: Row 20 cell 5 value is -ve or 0.00: {difference}, replacing by 0 instead.")
                table.rows[row_pos].cells[4].text = rupee_symbol + string_zero
        elif sgst1 is None:
            table.rows[row_pos].cells[4].text = gstr3b_analysis_NA
            print(
                f"Row 20 cell 5: Either one or both of the SGST values is not proper- sgst1: {sgst1}, sgst2 = {sgst2}")
        elif sgst2 is None:
            table.rows[row_pos].cells[4].text = gstr2a_analysis_NA
            print(
                f"Row 20 cell 5: Either one or both of the SGST values is not proper- sgst1: {sgst1}, sgst2 = {sgst2}")

        if isinstance(cess1, numbers.Number) and isinstance(cess2, numbers.Number):
            difference = cess1 - cess2
            if difference > 0:  # populate if value is +ve
                table.rows[row_pos].cells[5].text = rupee_symbol + str(difference)
            else:
                print(f"Not accounted: Row 20 cell 6 value is -ve or 0.00: {difference}, replacing by 0 instead.")
                table.rows[row_pos].cells[5].text = rupee_symbol + string_zero
        elif cess1 is None:
            table.rows[row_pos].cells[5].text = gstr3b_analysis_NA
            print(
                f"Row 20 cell 6: Either one or both of the CESS values is not proper- cess1: {cess1}, cess2 = {cess2}")
        elif cess2 is None:
            table.rows[row_pos].cells[5].text = gstr2a_analysis_NA
            print(
                f"Row 20 cell 6: Either one or both of the CESS values is not proper- cess1: {cess1}, cess = {cess2}")
        row_pos += 1

        # 21. Row 21 cell 3
        igst1 = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_21_IGST', None)
        cgst1 = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_21_CGST', None)
        sgst1 = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_21_SGST', None)
        cess1 = master_dict.get('gstr3b_analysis_dict', {}).get('result_point_21_CESS', None)
        igst2 = master_dict.get('gstr2a_analysis_dict', {}).get('result_point_21_IGST', None)
        cgst2 = master_dict.get('gstr2a_analysis_dict', {}).get('result_point_21_CGST', None)
        sgst2 = master_dict.get('gstr2a_analysis_dict', {}).get('result_point_21_SGST', None)
        cess2 = master_dict.get('gstr2a_analysis_dict', {}).get('result_point_21_CESS', None)
        if isinstance(igst1, numbers.Number) and isinstance(igst2, numbers.Number):
            difference = igst1 - igst2
            if difference > 0:  # populate if value is +ve
                table.rows[row_pos].cells[2].text = rupee_symbol + str(difference)
            else:
                print(f"Not accounted: Row 21 cell 3 value is -ve or 0.00: {difference}, replacing by 0 instead.")
                table.rows[row_pos].cells[2].text = rupee_symbol + string_zero
        elif igst1 is None:
            table.rows[row_pos].cells[2].text = gstr3b_analysis_NA
            print(
                f"Row 21 cell 3: Either one or both of the IGST values is not proper- igst1: {igst1}, igst2 = {igst2}")
        elif igst2 is None:
            table.rows[row_pos].cells[2].text = gstr2a_analysis_NA
            print(
                f"Row 21 cell 3: Either one or both of the IGST values is not proper- igst1: {igst1}, igst2 = {igst2}")

        if isinstance(cgst1, numbers.Number) and isinstance(cgst2, numbers.Number):
            difference = cgst1 - cgst2
            if difference > 0:  # populate if value is +ve
                table.rows[row_pos].cells[3].text = rupee_symbol + str(difference)
            else:
                print(f"Not accounted: Row 21 cell 4 value is -ve or 0.00: {difference}, replacing by 0 instead.")
                table.rows[row_pos].cells[3].text = rupee_symbol + string_zero
        elif cgst1 is None:
            table.rows[row_pos].cells[3].text = gstr3b_analysis_NA
            print(
                f"Row 21 cell 4: Either one or both of the CGST values is not proper- cgst1: {cgst1}, cgst2 = {cgst2}")
        elif cgst2 is None:
            table.rows[row_pos].cells[3].text = gstr2a_analysis_NA
            print(
                f"Row 21 cell 4: Either one or both of the CGST values is not proper- cgst1: {cgst1}, cgst2 = {cgst2}")

        if isinstance(sgst1, numbers.Number) and isinstance(sgst2, numbers.Number):
            difference = sgst1 - sgst2
            if difference > 0:  # populate if value is +ve
                table.rows[row_pos].cells[4].text = rupee_symbol + str(difference)
            else:
                print(f"Not accounted: Row 21 cell 5 value is -ve or 0.00: {difference}, replacing by 0 instead.")
                table.rows[row_pos].cells[4].text = rupee_symbol + string_zero
        elif sgst1 is None:
            table.rows[row_pos].cells[4].text = gstr3b_analysis_NA
            print(
                f"Row 21 cell 5: Either one or both of the SGST values is not proper- sgst1: {sgst1}, sgst2 = {sgst2}")
        elif sgst2 is None:
            table.rows[row_pos].cells[4].text = gstr2a_analysis_NA
            print(
                f"Row 21 cell 5: Either one or both of the SGST values is not proper- sgst1: {sgst1}, sgst2 = {sgst2}")

        if isinstance(cess1, numbers.Number) and isinstance(cess2, numbers.Number):
            difference = cess1 - cess2
            if difference > 0:  # populate if value is +ve
                table.rows[row_pos].cells[5].text = rupee_symbol + str(difference)
            else:
                print(f"Not accounted: Row 21 cell 6 value is -ve or 0.00: {difference}, replacing by 0 instead.")
                table.rows[row_pos].cells[5].text = rupee_symbol + string_zero
        elif cess1 is None:
            table.rows[row_pos].cells[5].text = gstr3b_analysis_NA
            print(
                f"Row 21 cell 6: Either one or both of the CESS values is not proper- cess1: {cess1}, cess2 = {cess2}")
        elif cess2 is None:
            table.rows[row_pos].cells[5].text = gstr2a_analysis_NA
            print(
                f"Row 21 cell 6: Either one or both of the CESS values is not proper- cess1: {cess1}, cess = {cess2}")
        row_pos += 1

        # 22. Row 22 cell 3
        igst = master_dict.get('gstr3b_merged_dict', {}).get('result_point_22_IGST', None)
        cgst = master_dict.get('gstr3b_merged_dict', {}).get('result_point_22_CGST', None)
        sgst = master_dict.get('gstr3b_merged_dict', {}).get('result_point_22_SGST', None)
        cess = master_dict.get('gstr3b_merged_dict', {}).get('result_point_22_CESS', None)
        if igst is not None:
            table.rows[row_pos].cells[2].text = rupee_symbol + str(igst)
        else:
            table.rows[row_pos].cells[2].text = gstr3b_merged_NA
        if cgst is not None:
            table.rows[row_pos].cells[3].text = rupee_symbol + str(cgst)
        else:
            table.rows[row_pos].cells[3].text = gstr3b_merged_NA
        if sgst is not None:
            table.rows[row_pos].cells[4].text = rupee_symbol + str(sgst)
        else:
            table.rows[row_pos].cells[4].text = gstr3b_merged_NA
        if cess is not None:
            table.rows[row_pos].cells[5].text = rupee_symbol + str(cess)
        else:
            table.rows[row_pos].cells[5].text = gstr3b_merged_NA
        row_pos += 1

        # Save the document
        doc.save(output_path)
        print(f"✅ [General report generator] Document saved at: {output_path}")
    except Exception as e:
        print(f"[General report generator] ❌ Error during report generation: {e}")

